from flask import Flask, request, jsonify, render_template, send_file, redirect, url_for
import os
import openai
from openai import OpenAI
import google.generativeai as genai
import chromadb
from langchain_openai import OpenAIEmbeddings
import psycopg2
import json
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.shared
from docx.oxml import parse_xml
from bs4 import BeautifulSoup
import random
# ==== 登入 ====
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
# ==== Email驗證 ====
import smtplib
from email.mime.text import MIMEText
import string
from datetime import datetime, timezone, timedelta
import pytesseract
from PIL import Image
import io
import base64
from pdf2image import convert_from_bytes
# ==== Selenium ====
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

# ==== Tesseract 設定 ====
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'  # 容器中的路徑
poppler_path = None  # 在 Linux 容器中不需要指定 poppler 路徑

# ==== Flask 設定 ====
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "a-very-secret-dev-key")

# ==== 登入管理設定 ====
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'  # 如果未登入會跳到 /login

class User(UserMixin):
    """用戶類別"""
    def __init__(self, id, email, plan):
        self.id = id
        self.email = email
        self.plan = plan

@login_manager.user_loader
def load_user(user_id):
    """載入用戶"""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT id, email, plan FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        if user:
            return User(str(user[0]), user[1], user[2])
        return None
    finally:
        cursor.close()
        conn.close()

@app.before_request
def require_login():
    allowed_routes = ['login', 'register', 'verify', 'static', 'index']  # 不用登入也能看的路由
    if request.endpoint not in allowed_routes and not current_user.is_authenticated:
        return redirect(url_for('login'))
    
# ==== 註冊 ====
# 寄送驗證碼函數
def send_verification_email(to_email, code):
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    sender_email = os.getenv("SMTP_EMAIL")
    sender_password = os.getenv("SMTP_PASSWORD")

    # 檢查必要的環境變數
    if not sender_email or not sender_password:
        print("❌ 缺少 SMTP 設定：請確認 SMTP_EMAIL 和 SMTP_PASSWORD 環境變數已設置")
        raise ValueError("Missing SMTP credentials")

    # 準備郵件內容
    msg = MIMEText(f"你的備審系統註冊驗證碼是：{code}，請於10分鐘內完成驗證。")
    msg['Subject'] = '【備審系統】Email驗證碼'
    msg['From'] = sender_email
    msg['To'] = to_email

    try:
        # 建立安全連接
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.ehlo()  # 可以幫助識別連接問題
        server.starttls()
        server.ehlo()  # 重新識別

        try:
            # 使用顯式的登入方法
            server.login(sender_email, sender_password)
            print(f"✅ SMTP 登入成功")
            
            # 發送郵件
            server.sendmail(sender_email, to_email, msg.as_string())
            print(f"✅ 驗證信已發送至 {to_email}")
            
        except smtplib.SMTPAuthenticationError as auth_error:
            print(f"❌ SMTP 認證失敗: {auth_error}")
            raise ValueError(f"SMTP authentication failed: {str(auth_error)}")
            
        except smtplib.SMTPException as smtp_error:
            print(f"❌ SMTP 錯誤: {smtp_error}")
            raise ValueError(f"SMTP error: {str(smtp_error)}")
            
        finally:
            try:
                server.quit()
            except Exception as e:
                print(f"⚠️ 關閉 SMTP 連接時發生錯誤: {e}")
                
    except Exception as e:
        print(f"❌ SMTP 連接失敗: {str(e)}")
        raise ValueError(f"Failed to connect to SMTP server: {str(e)}")

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        try:
            email = request.form['email']
            password = request.form['password']
            confirm_password = request.form['confirm_password']

            if password != confirm_password:
                return jsonify({
                    "success": False,
                    "message": "❌ 密碼不一致"
                }), 400

            conn = get_db_connection()
            cursor = conn.cursor()

            try:
                # 檢查 email 是否存在
                cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
                existing_user = cursor.fetchone()
                if existing_user:
                    return jsonify({
                        "success": False,
                        "message": "❌ Email 已註冊"
                    }), 400

                # 產生隨機驗證碼
                code = ''.join(random.choices(string.digits, k=6))
                hashed_password = generate_password_hash(password)
                now = datetime.now(timezone.utc)

                # 先儲存用戶資料
                cursor.execute(
                    "INSERT INTO users (email, password, plan, is_verified, verification_code, verification_sent_at) VALUES (%s, %s, %s, %s, %s, %s)",
                    (email, hashed_password, 'Free', False, code, now)
                )
                conn.commit()

                try:
                    # 嘗試發送驗證郵件
                    send_verification_email(email, code)
                    return jsonify({
                        "success": True,
                        "redirect": url_for('verify', email=email)
                    })
                    
                except Exception as e:
                    # 如果發送失敗，刪除剛才新增的用戶資料
                    cursor.execute("DELETE FROM users WHERE email = %s", (email,))
                    conn.commit()
                    print(f"❌ 驗證信發送失敗，已刪除用戶資料: {str(e)}")
                    return jsonify({
                        "success": False,
                        "message": "❌ 驗證信發送失敗，請稍後再試或聯繫管理員"
                    }), 500

            except Exception as e:
                conn.rollback()
                print(f"❌ 註冊過程發生錯誤: {str(e)}")
                return jsonify({
                    "success": False,
                    "message": "❌ 註冊失敗，請稍後再試"
                }), 500
            finally:
                cursor.close()
                conn.close()

        except Exception as e:
            print(f"❌ 處理註冊請求時發生錯誤: {str(e)}")
            return jsonify({
                "success": False,
                "message": "❌ 系統錯誤，請稍後再試"
            }), 500

    return render_template('register.html')

@app.route('/verify', methods=['GET', 'POST'])
def verify():
    email = request.args.get('email')
    if not email:
        return redirect(url_for('register'))

    error_message = None

    if request.method == 'POST':
        input_code = request.form['verification_code']

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT verification_code, verification_sent_at FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()

        if not user:
            cursor.close()
            conn.close()
            error_message = "查無此Email，請重新註冊"
        else:
            code_from_db, sent_time = user
            now = datetime.utcnow()

            # 驗證碼是否過期（有效10分鐘）
            if sent_time is None or now - sent_time > timedelta(minutes=10):
                error_message = "驗證碼已過期，請重新寄送驗證碼。"
            elif input_code == code_from_db:
                cursor.execute(
                    "UPDATE users SET is_verified = TRUE, verification_code = NULL, verification_sent_at = NULL WHERE email = %s",
                    (email,)
                )
                conn.commit()
                cursor.close()
                conn.close()
                return redirect(url_for('login'))
            else:
                error_message = "驗證碼錯誤，請重新輸入或重新寄送驗證碼。"
            
            cursor.close()
            conn.close()

    return render_template('verify.html', email=email, error_message=error_message)

@app.route('/resend_verification', methods=['POST'])
def resend_verification():
    email = request.form['email']

    conn = get_db_connection()
    cursor = conn.cursor()

    new_code = ''.join(random.choices(string.digits, k=6))
    now = datetime.utcnow()

    cursor.execute(
        "UPDATE users SET verification_code = %s, verification_sent_at = %s WHERE email = %s",
        (new_code, now, email)
    )
    conn.commit()
    cursor.close()
    conn.close()

    send_verification_email(email, new_code)

    return redirect(url_for('verify', email=email))

# ==== Profile ====
@app.route("/profile")
@login_required
def profile():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 獲取用戶基本信息
    cursor.execute("""
        SELECT email, plan, is_verified, generation_quota, generation_count, plan_expiration 
        FROM users 
        WHERE id = %s
    """, (current_user.id,))
    user = cursor.fetchone()
    
    # 獲取生成紀錄
    cursor.execute("""
        SELECT id, university, department, section_code, generated_text, 
               style, created_at, is_favorite
        FROM generation_history 
        WHERE user_id = %s 
        ORDER BY created_at DESC
    """, (current_user.id,))
    history = cursor.fetchall()
    
    cursor.close()
    conn.close()

    if not user:
        return "找不到使用者", 404

    user_info = {
        "email": user[0],
        "plan": user[1],
        "is_verified": user[2],
        "generation_quota": user[3],
        "generation_count": user[4],
        "plan_expiration": user[5]
    }
    
    # 將查詢結果轉換為字典列表
    history_list = []
    for record in history:
        history_list.append({
            "id": record[0],
            "university": record[1],
            "department": record[2],
            "section_code": record[3],
            "generated_text": record[4],
            "style": record[5],
            "created_at": record[6],
            "is_favorite": record[7]
        })

    return render_template("profile.html", user=user_info, history=history_list)

# ==== 載入 .env ====
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")

# 初始化 OpenAI
openai.api_key = OPENAI_API_KEY
openai_client = openai.OpenAI()

# 初始化 Gemini
genai.configure(api_key=GEMINI_API_KEY)

# ==== 向量資料庫 (備用學系特色檢索) ====
try:
    embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
    db = chromadb.PersistentClient(path="./chroma_db")
    collection = db.get_or_create_collection(
        name="department_info",
        metadata={"hnsw:space": "cosine"}
    )
except Exception as e:
    print(f"Warning: ChromaDB initialization failed: {str(e)}")
    print("The system will continue without vector database support.")
    db = None
    collection = None

# ==== 連接 PostgreSQL 資料庫 ====
def get_db_connection():
    return psycopg2.connect(DATABASE_URL)

# ==== 產生亂碼ID工具 ====
def generate_random_id():
    """生成6位亂碼ID，包含數字和大寫字母"""
    characters = string.ascii_uppercase + string.digits  # A-Z and 0-9
    while True:
        # 生成6位亂碼
        random_id = ''.join(random.choices(characters, k=6))
        # 檢查ID是否已存在
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM generation_history WHERE id = %s", (random_id,))
        exists = cursor.fetchone()
        cursor.close()
        conn.close()
        # 如果ID不存在，則返回這個ID
        if not exists:
            return random_id

# ==== 代碼對應表 ====
CODE_MAPPING = {
    "A": "修課紀錄", "B": "書面報告", "C": "實作作品", "D": "自然科學探究",
    "E": "社會領域探究", "F": "高中自主學習計畫", "G": "社團活動經驗", "H": "幹部經驗",
    "I": "服務學習經驗", "J": "競賽表現", "K": "非修課成果作品", "L": "檢定證照",
    "M": "特殊優良表現", "N": "多元表現綜整心得", "O": "高中學習歷程反思",
    "P": "學習動機", "Q": "未來學習計畫"
}

# ==== 分項表與Prompt ====
SECTION_MAPPING = {
    "ABC": {  
        "codes": ["A", "B", "C", "D", "E"],
        "prompt": "請撰寫一篇課程學習成果，強調學習歷程與探索精神，並與學系關聯性相結合。"
    },
    "N": {  
        "codes": ["F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請了解學系專業與選材理念後，撰寫強調自主學習的多元表現綜整心得，包含引言、3個段落(其中每個段落需有一個標題，總結該段落的內容)，以及總結，並利用使用者輸入的具體事例舉例每件事的能力成長與省思。內容與學系密切相關，且必須確保字數在800字內。。"
    },
    "O": {  
        "codes": ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請撰寫高中學習歷程反思，強調學習經驗、挑戰與成長，並舉出失敗與反省以展現學習態度。必須確保字數在800字內。"
    },
    "P": {  
        "codes": ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請撰寫學習動機，強調與學系高度相關的興趣與企圖心，並適當佐證。",
        "needs_department_features": True
    },
    "Q": {  
        "codes": ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"],
        "prompt": "請撰寫未來計畫，分成三個階段：(1)入學前想學習的課程與技能 (2)就讀期間計畫修習的課程 (3)畢業後的職業發展。",
        "needs_department_features": True
    }
}

# ==== 風格Prompt ====
STYLE_PROMPTS = {
    "formal": "句子可以稍口語、真誠、有情緒，但仍保持基本邏輯與清晰度。避免誇張詞彙與過度文學。",
    "casual": "請以自然、口語化的語氣撰寫，像是在向教授講述自己的故事。內容可稍微白話，但仍應保持邏輯與專業度。",
    "concise": "請以簡潔明快的語言撰寫，避免冗詞贅句，直接切入重點，強調邏輯清晰與重點聚焦。",
    "detailed": "請以完整、深入的方式撰寫，內容須具體，包含明確例子與說明，讓讀者能全面理解申請者的能力與動機。",
}

# ====登入登出====
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        if not email or not password:
            return jsonify({
                "success": False,
                "message": "請輸入電子郵件和密碼",
                "code": "MISSING_FIELDS"
            })

        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT id, email, password, plan, is_verified FROM users WHERE email = %s", (email,))
            user = cursor.fetchone()
            cursor.close()
            conn.close()

            if not user:
                return jsonify({
                    "success": False,
                    "message": "此電子郵件尚未註冊",
                    "code": "NOT_REGISTERED"
                })

            if not check_password_hash(user[2], password):
                return jsonify({
                    "success": False,
                    "message": "密碼錯誤",
                    "code": "WRONG_PASSWORD"
                })

            if not user[4]:  # is_verified
                return jsonify({
                    "success": False,
                    "message": "帳號尚未完成 Email 驗證，請先驗證後登入",
                    "code": "NOT_VERIFIED"
                })

            user_obj = User(id=user[0], email=user[1], plan=user[3])
            login_user(user_obj)
            
            return jsonify({
                "success": True,
                "message": "登入成功",
                "redirect": url_for('index')
            })

        except Exception as e:
            print(f"登入錯誤: {str(e)}")
            return jsonify({
                "success": False,
                "message": "登入時發生錯誤，請稍後再試",
                "code": "SERVER_ERROR"
            })

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# ==== Check Quota ====
@app.route("/check_quota")
@login_required
def check_quota():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 檢查用戶的方案和配額
        cursor.execute("""
            SELECT plan, generation_quota, generation_count, plan_expiration
            FROM users 
            WHERE id = %s
        """, (current_user.id,))
        
        result = cursor.fetchone()
        if not result:
            return jsonify({"error": "找不到用戶資料"}), 404
            
        plan, quota, count, expiration = result
        
        # 如果是無限方案，返回-1表示無限制
        if plan == "Unlimited":
            remaining_quota = -1
        else:
            remaining_quota = quota - count
            
        # 檢查方案是否過期
        is_expired = False
        if expiration and expiration < datetime.now().date():
            is_expired = True
            
        return jsonify({
            "plan": plan,
            "total_quota": quota,
            "used_quota": count,
            "remaining_quota": remaining_quota,
            "is_expired": is_expired,
            "expiration_date": expiration.strftime('%Y-%m-%d') if expiration else None
        })
        
    except Exception as e:
        print(f"檢查配額時發生錯誤: {str(e)}")
        return jsonify({"error": "檢查配額時發生錯誤"}), 500
    finally:
        cursor.close()
        conn.close()

# ==== Purchase ====
@app.route("/purchase", methods=["GET", "POST"])
@login_required
def purchase():
    if request.method == "POST":
        # 模擬購買成功
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE users SET plan = %s WHERE id = %s", ("Premium", current_user.id))
        conn.commit()
        cursor.close()
        conn.close()

        return redirect(url_for("profile"))

    # 預設 GET 顯示畫面
    return render_template("purchase.html", current_plan=current_user.plan)

# ==== QA ====
@app.route("/ask", methods=["POST"])
def ask_question():
    data = request.get_json()
    university = data.get("university")
    department = data.get("department")
    question = data.get("question")

    if not university or not department or not question:
        return jsonify({"answer": "請提供完整的學校、學系與問題內容！"}), 400

    # 取得該學系特色與需繳交資料作為回答基礎
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT section_code, section_name, content 
        FROM application_guidelines 
        WHERE university = %s AND department = %s
    """, (university, department))
    sections = cursor.fetchall()
    cursor.close()
    conn.close()

    department_features = get_department_features(university, department)

    base_context = "\n".join([
        f"{code} - {name}：{content}" for code, name, content in sections
    ]) or "查無備審項目內容。"

    prompt = f"""
你是一名大學備審客服人員，請針對使用者提出的問題，根據下列內容進行回答：
- **請務必將每一條條列點獨立寫在新的一行（請用 \n 分隔），不要將多個條列點寫在同一行。**


【學校】{university}
【學系】{department}
【系所特色】{department_features}
【備審資料】{base_context}

使用者提問：
{question}

請以簡潔清楚的語氣回覆，字數不超過150字，必要時可引用學系特色或備審資料，但不得亂編。
    """

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        answer = response.choices[0].message.content.strip()
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"answer": f"⚠️ 回覆錯誤：{str(e)}"}), 500



# ==== 提取學系特色 (優先從資料庫，其次從向量資料庫) ====
def get_department_features(university, department):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT features, suggestion_content FROM department_features 
        WHERE university = %s AND department = %s
    """, (university, department))
    result = cursor.fetchone()
    
    cursor.close()
    conn.close()

    if result:
        return {
            "features": result[0] if result[0] else "未提供學系特色",
            "suggestions": result[1] if result[1] else "未提供備審建議"
        }
    
    return {
        "features": "未提供學系特色",
        "suggestions": "未提供備審建議"
    }

# ==== 查詢需繳交資料與學系特色（首頁） ====
@app.route("/check", methods=["GET", "POST"])
def check_requirements():
    results = None
    features = None
    suggestions = None
    university = department = ""

    if request.method == "POST":
        university = request.form.get("university")
        department = request.form.get("department")

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT section_code, section_name FROM application_guidelines WHERE university=%s AND department=%s", (university, department))
        results = cursor.fetchall()
        
        # 獲取系所特色和備審建議
        dept_info = get_department_features(university, department)
        features = dept_info["features"]
        suggestions = dept_info["suggestions"]
        
        cursor.close()
        conn.close()

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT university, department FROM application_guidelines")
    all_options = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template(
        "check_requirements.html",
        all_options=all_options,
        results=results,
        features=features,
        suggestions=suggestions,
        university=university,
        department=department
    )

# ==== 獲取所有學校與學系(check) ====
@app.route("/get_departments", methods=["GET"])
def get_departments_by_school():
    university = request.args.get("university")
    if not university:
        return jsonify([])

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT DISTINCT department 
        FROM application_guidelines 
        WHERE university = %s
    """, (university,))
    departments = [row[0] for row in cursor.fetchall()]
    cursor.close()
    conn.close()
    
    return jsonify(departments)

# ==== 獲取所有學校與學系(generator) ====
@app.route("/get_universities", methods=["GET"])
def get_universities():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT university FROM application_guidelines")
    universities = [row[0] for row in cursor.fetchall()]
    cursor.close()
    conn.close()
    return jsonify(universities)


# ==== 學系備審指引查詢 ====
@app.route("/get_guidelines", methods=["POST"])
def get_guidelines():
    data = request.json
    university = data.get("university")
    department = data.get("department")

    if not university or not department:
        return jsonify({"error": "請選擇學校與學系"}), 400

    conn = get_db_connection()
    cursor = conn.cursor()

    # 查詢 application_guidelines
    cursor.execute("""
        SELECT section_code, section_name, content 
        FROM application_guidelines 
        WHERE university=%s AND department=%s
    """, (university, department))
    results = cursor.fetchall()

    # 查詢 department_features
    department_features = get_department_features(university, department)

    cursor.close()
    conn.close()

    if not results:
        return jsonify({"error": "該學校學系尚無備審資料"}), 400

    sections = {row[0]: {"name": row[1], "content": row[2]} for row in results}

    return jsonify({
        "university": university,
        "department": department,
        "sections": sections,
        "department_features": department_features
    })

@app.route("/check_auth")
def check_auth():
    """檢查用戶是否已登入"""
    if current_user.is_authenticated:
        return jsonify({"authenticated": True})
    return jsonify({"authenticated": False})

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["GET"])
def generate_page():
    # 從 URL 參數獲取學校和學系
    university = request.args.get('university')
    department = request.args.get('department')
    return render_template("generator.html", university=university, department=department)

# ==== 查詢備審指引頁（如果有）====
@app.route("/show_check_page", methods=["GET", "POST"])
def show_check_page():
    # 處理查詢資料邏輯
    return render_template("check_requirements.html")

#清除單字「分」、空段落、破碎文字段落
def clean_paragraph_output(text):
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join([line for line in lines if line and line != "分"])
    return cleaned

# ==== 生成段落 ====
@app.route("/generate_paragraph", methods=["POST"])
def generate_paragraph():
    # === 檢查剩餘次數 ===
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT generation_quota FROM users WHERE id = %s", (current_user.id,))
    result = cursor.fetchone()
    cursor.close()
    conn.close()

    if result and result[0] != -1 and result[0] <= 0:
        return jsonify({"error": "❗ 生成次數已用完，請升級為 Premium 會員以獲得更多次數！"}), 403
    
    try:
        data = request.get_json()
        
        section_code = data.get("section_code")
        user_inputs = data.get("user_inputs", {})
        university = data.get("university", "未知大學")
        department = data.get("department", "未知學系")
        style = data.get("style", "formal")  
        word_count = random.randint(600, 800) 
        adjust_percentage = random.randint(30, 50) 
        
        if section_code not in SECTION_MAPPING:
            return jsonify({"error": "無效的段落代碼"}), 400
        
        # 獲取學系特色
        department_features = get_department_features(university, department)
            
        # 獲取撰寫指引 (section_prompt)
        section_prompt = SECTION_MAPPING[section_code]["prompt"]
        
        input_text = "\n".join(f"{CODE_MAPPING.get(code, code)}: {text}" for code, text in user_inputs.items() if text)

        # 使用LLM初次生成，加入表格生成的指示
        first_prompt = f"""
    你是一名高中生，正在申請{university}-{department}，請根據以下內容撰寫 {section_prompt}。

    請務必遵守以下規則，嚴格按照使用者提供的內容進行撰寫，不得編造資訊：
    1.只能使用以下學系特色，不得新增額外內容：{department_features}
    2.只能使用使用者提供的內容，不得自行發揮：{input_text}
    3.字數範圍：{word_count} 字以內
    4.風格要求：{STYLE_PROMPTS.get(style, STYLE_PROMPTS['formal'])}
    5.不得杜撰經歷、不得添加未提供的競賽、活動、學術研究等內容。
    6.內容應清晰、邏輯合理、段落流暢，並忠實呈現使用者輸入的重點。

    特別要求：
    1.請判斷內容是否適合使用表格呈現（例如：競賽成果、證照、課程學習、社團活動等）
    2.表格應包含以下元素：
       - 基本事實（如：活動名稱、時間、成果等）
       - 學習心得（如：遇到的挑戰、解決方法）
       - 能力反思（如：培養了什麼能力、對申請學系有何幫助）
    3.表格結構要求：
       - 使用HTML格式，包含<table>、<tr>、<td>等標籤
       - 表格標題應反映內容主題
       - 至少包含「事實描述」、「學習心得」、「能力反思」等欄位
       - 可根據內容性質自行調整或增加欄位
    4.表格呈現方式：
       - 每個重要經歷都應包含完整的反思內容
       - 心得和反思應具體且深入，不要流於表面
       - 確保反思內容與學系特色相呼應
    5.表格應放在適當的段落位置，並與文字內容自然銜接
    6.可以根據內容性質設計多個不同主題的表格

    請根據這些要求，產生一段符合申請需求的內容。
    """
        # 使用 Gemini 進行初次生成
        try:
            gemini_model = genai.GenerativeModel("gemini-2.0-flash")
            gemini_response = gemini_model.generate_content(
                first_prompt, 
                generation_config={
                    "temperature": 0.6,
                    "max_output_tokens": 2048,
                    "top_p": 0.8,
                    "top_k": 40
                }
            )
            first_output = gemini_response.text.strip()
        except Exception as e:
            print(f"Gemini 生成失敗: {str(e)}")
            return jsonify({"error": f"內容生成失敗: {str(e)}"}), 500

        # 進行優化，保留表格結構
        try:
            improved_prompt = f"""
    你是一位擅長潤飾大學申請備審資料的語言專家，請協助我修改以下備審資料，使其自然流暢，語氣真誠，並保留原有內容重點。

    目前的內容：
    {first_output}

    請確保：
    1.不可新增資訊，只能在原始內容基礎上進行潤飾與語句調整
    2.風格請採用：{STYLE_PROMPTS.get(style, STYLE_PROMPTS['formal'])}
    3.請避免使用以下 AI 常見用語或句型：如「總體而言」、「本文將探討」、「綜上所述」、「在當今社會中」、「產生深遠影響」、「我堅信」等
    4.請避免過於工整、生硬、過度學術化的句式結構，讓整體語氣更貼近一位真誠且有思考力的高中學生口吻
    5.每個新段落請用 \n\n 分隔，輸出為純文字，不使用 Markdown 語法
    6.請將文字濃縮至 {word_count} 字以內，保留關鍵資訊與主要邏輯，刪除重複或不必要的詞語，使內容更精練但不失自然語感
    7.優化表格內容：
       - 確保心得和反思的內容真實具體，避免空泛的描述
       - 調整表格中的文字，使其更簡潔有力
       - 保持反思內容的深度和個人特色
    8.確保表格與周圍文字的自然銜接，使整體內容更有層次感和邏輯性

    優化後的內容：
    """
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": improved_prompt}],
                temperature=0.5,
                timeout=180,  # 設置超時時間為 180 秒
            )
            improved_output = response.choices[0].message.content.strip()
            
            # === 儲存生成記錄 ===
            conn = get_db_connection()
            cursor = conn.cursor()
            history_id = generate_random_id()
            cursor.execute("""
                INSERT INTO generation_history 
                (id, user_id, university, department, section_code, generated_text, mindmap_data, style)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                history_id, current_user.id, university, department, section_code,
                improved_output, json.dumps(data.get('mindmap_data', {})), style
            ))
            
            # === 扣除配额 ===
            if result and result[0] != -1:
                cursor.execute(
                    "UPDATE users SET generation_quota = generation_quota - 1 WHERE id = %s",
                    (current_user.id,)
                )
            conn.commit()
            cursor.close()
            conn.close()

            return jsonify({
                "style": style,
                "adjust_percentage": adjust_percentage,
                "generated_text": improved_output,
                "history_id": history_id
            }), 200, {'Content-Type': 'application/json; charset=utf-8'}
            
        except Exception as e:
            print(f"OpenAI 生成失敗: {str(e)}")
            return jsonify({"error": f"內容優化失敗: {str(e)}"}), 500
            
    except Exception as e:
        print(f"生成段落時發生錯誤: {str(e)}")
        return jsonify({"error": f"處理請求失敗: {str(e)}"}), 500

# ====重新部分生成====
@app.route("/regenerate_paragraph", methods=["POST"])
def regenerate_paragraph():
    data = request.get_json()
    original_paragraph = data.get("original_paragraph")
    style = data.get("style", "formal")

    if not original_paragraph:
        return jsonify({"error": "缺少原始段落"}), 400

    refine_prompt = f"""
    這是要重新優化的段落：
    {original_paragraph}

    請確保：
    1.請調整語句，使表達方式稍有不同，但仍然保持相同核心內容
    2.風格請採用：{STYLE_PROMPTS.get(style, STYLE_PROMPTS['formal'])}
    3.請避免使用以下 AI 常見用語或句型：如「總體而言」、「本文將探討」、「綜上所述」、「在當今社會中」、「產生深遠影響」、「我堅信」等
    4.請避免過於工整、生硬、過度學術化的句式結構，讓整體語氣更貼近一位真誠且有思考力的高中學生口吻
    5.每個新段落請用 \n\n 分隔，輸出為純文字，不使用 Markdown 語法
    6.請將文字濃縮至 150 字以內，保留關鍵資訊與主要邏輯，刪除重複或不必要的詞語，使內容更精練但不失自然語感。
    """

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": refine_prompt}],
            temperature=0.6,
        )
        new_paragraph = response.choices[0].message.content.strip()
        return jsonify({"new_paragraph": new_paragraph})
    except Exception as e:
        return jsonify({"error": f"重新生成時發生錯誤: {str(e)}"}), 500


# ==== 生成 Word 文件 ====
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    data = request.get_json()
    university = data.get("university", "未知大學")
    department = data.get("department", "未知學系")
    section_code = data.get("section_code", "未知段落")
    section_name = CODE_MAPPING.get(section_code, section_code)
    generated_text = data.get("generated_text", "")
    mindmap = data.get("mindmap", "")  # 獲取心智圖數據
    
    if not generated_text:
        return jsonify({"error": "沒有生成內容"}), 400

    doc = Document()
    
    # 設置文件樣式
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style.font.size = docx.shared.Pt(12)
    
    # 添加標題
    title = doc.add_heading(f"{university}-{department}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f"{section_code}-{section_name}", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 如果有心智圖，先添加心智圖部分
    if mindmap:
        doc.add_heading("申請重點分析", level=3)
        doc.add_paragraph("以下心智圖展示了申請者的經歷如何符合學系需求：")
        
        # 將 Mermaid 語法轉換為文字形式的心智圖
        try:
            # 解析 Mermaid 語法並轉換為文字大綱
            lines = mindmap.split('\n')
            current_level = 0
            for line in lines:
                if line.strip() and not line.strip().startswith('mindmap'):
                    # 計算縮排級別
                    indent_level = (len(line) - len(line.lstrip())) // 2
                    content = line.strip()
                    
                    # 移除 Mermaid 語法中的特殊字符
                    content = content.replace('((', '').replace('))', '')
                    content = content.replace('[', '').replace(']', '')
                    
                    # 根據縮排級別設置不同的格式
                    if indent_level == 0:
                        p = doc.add_paragraph()
                        run = p.add_run(content)
                        run.bold = True
                        run.font.size = docx.shared.Pt(14)
                    else:
                        # 添加縮排和項目符號
                        p = doc.add_paragraph()
                        p.style = 'List Bullet'
                        p.paragraph_format.left_indent = docx.shared.Pt(36 * indent_level)
                        run = p.add_run(content)
                        run.font.size = docx.shared.Pt(12)
            
            # 添加分隔
            doc.add_paragraph()
            
        except Exception as e:
            print(f"心智圖轉換失敗: {str(e)}")
    
    # 添加生成的內容
    doc.add_heading("申請內容", level=3)
    
    # 利用 BeautifulSoup 解析 HTML 結構
    soup = BeautifulSoup(generated_text, "html.parser")
    
    for element in soup.contents:
        if element.name == "h4":
            # 添加小標題
            doc.add_heading(element.text.strip(), level=4)
        
        elif element.name == "table":
            # 處理表格
            rows = element.find_all("tr")
            if rows:
                # 獲取表格標題（如果有的話）
                table_title = None
                prev_element = element.find_previous_sibling()
                if prev_element and prev_element.name in ['h4', 'h5', 'h6']:
                    table_title = prev_element.text.strip()
                
                if table_title:
                    title_paragraph = doc.add_paragraph()
                    title_run = title_paragraph.add_run(table_title)
                    title_run.bold = True
                    title_run.font.size = docx.shared.Pt(12)
                
                # 創建表格
                num_cols = len(rows[0].find_all(["td", "th"]))
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Table Grid'
                
                # 設置表格寬度
                table.autofit = False
                table_width = docx.shared.Inches(6.0)
                table.width = table_width
                
                # 處理每一行
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    new_row = table.add_row().cells
                    
                    # 計算每列的寬度（根據內容長度）
                    col_widths = []
                    for cell in cells:
                        text_length = len(cell.text.strip())
                        col_widths.append(max(text_length * 0.2, 1))  # 最小寬度為1
                    
                    total_width = sum(col_widths)
                    for j, (cell, width) in enumerate(zip(cells, col_widths)):
                        cell_text = cell.text.strip()
                        new_row[j].text = cell_text
                        
                        # 設置單元格寬度
                        cell_width = int(table_width * (width / total_width))
                        new_row[j].width = cell_width
                        
                        # 如果是表頭，設置粗體
                        if i == 0:
                            paragraph = new_row[j].paragraphs[0]
                            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                            run.bold = True
                            run.font.size = docx.shared.Pt(11)
                            # 設置表頭背景色為淺灰色
                            new_row[j]._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F2F2F2"/>'))
                        
                        # 設置單元格文字對齊方式
                        paragraph = new_row[j].paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 在表格後添加空行
                doc.add_paragraph()
        
        elif element.name in ["p", "div"]:
            if element.text.strip():
                p = doc.add_paragraph(element.text.strip())
                # 設置段落間距
                p.paragraph_format.space_after = docx.shared.Pt(12)
        
        elif isinstance(element, str):
            if element.strip():
                p = doc.add_paragraph(element.strip())
                p.paragraph_format.space_after = docx.shared.Pt(12)
    
    # 使用 BytesIO 在記憶體中保存文件
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    # 生成檔案名稱
    filename = f"{university}{department}_{section_code}.docx"
    
    # 直接返回文件流
    return send_file(
        docx_file,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

def convert_mermaid_to_image(mermaid_code):
    """
    將 Mermaid 代碼轉換為圖片
    """
    try:
        # 建立臨時文件目錄
        temp_dir = os.path.join("static", "temp")
        os.makedirs(temp_dir, exist_ok=True)
        
        # 生成唯一的檔案名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_html = os.path.join(temp_dir, f"mindmap_{timestamp}.html")
        output_image = os.path.join(temp_dir, f"mindmap_{timestamp}.png")
        
        # 創建包含 Mermaid 的 HTML 文件，添加自定義樣式
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
            <style>
                .mermaid {{
                    background-color: white;
                }}
                /* 確保所有文字為黑色 */
                .mindmap-node .nodeText,
                .mindmap-node text,
                .mindmap-node .nodeLabel {{
                    color: black !important;
                    fill: black !important;
                }}
            </style>
            <script>
                mermaid.initialize({{
                    startOnLoad: true,
                    theme: 'default',
                    mindmap: {{
                        padding: 20,
                        useMaxWidth: true
                    }},
                    themeVariables: {{
                        // 設置所有文字顏色為黑色
                        primaryTextColor: '#000000',
                        secondaryTextColor: '#000000',
                        tertiaryTextColor: '#000000',
                        // 確保節點文字為黑色
                        nodeBorder: '#000000',
                        nodeTextColor: '#000000',
                        // 其他顏色設置
                        mainBkg: '#ffffff',
                        nodeBkg: '#ffffff'
                    }}
                }});
            </script>
        </head>
        <body>
            <div class="mermaid">
                {mermaid_code}
            </div>
        </body>
        </html>
        """
        
        with open(temp_html, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # 使用 Selenium 將 HTML 轉換為圖片
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        
        # 使用 webdriver_manager 自動管理 ChromeDriver
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        try:
            driver.get(f"file:///{os.path.abspath(temp_html)}")
            # 等待 Mermaid 渲染完成
            import time
            time.sleep(2)
            
            # 找到 Mermaid 圖表元素並截圖
            element = driver.find_element(By.CLASS_NAME, 'mermaid')
            element.screenshot(output_image)
            
            return output_image
            
        finally:
            driver.quit()
            # 清理臨時 HTML 文件
            if os.path.exists(temp_html):
                os.remove(temp_html)
            
    except Exception as e:
        print(f"轉換 Mermaid 圖表時發生錯誤: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return None

# ==== OCR 文件處理 ====
@app.route("/process_ocr", methods=["POST"])
@login_required
def process_ocr():
    if 'file' not in request.files or 'code' not in request.form:
        return jsonify({"error": "缺少必要參數"}), 400
    
    file = request.files['file']
    code = request.form['code']
    is_deleted = request.form.get('is_deleted', 'false').lower() == 'true'
    
    if is_deleted:
        return jsonify({
            "success": True,
            "text": "",
            "is_raw": False,
            "was_condensed": False,
            "skipped": True
        })
    
    if file.filename == '':
        return jsonify({"error": "沒有選擇檔案"}), 400
    
    try:
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension == '.pdf':
            # 處理 PDF 檔案
            pdf_bytes = file.read()
            try:
                # 在 Linux 容器中直接使用 convert_from_bytes
                images = convert_from_bytes(pdf_bytes)
                
                text_results = []
                for image in images:
                    try:
                        # 指定語言包，包含繁體中文和英文
                        text = pytesseract.image_to_string(
                            image, 
                            lang='chi_tra+eng',
                            config='--psm 1 --oem 1'  # 使用更準確的辨識模式
                        )
                        if text.strip():
                            text_results.append(text.strip())
                    except Exception as e:
                        print(f"OCR 處理單頁時發生錯誤: {str(e)}")
                        continue
                
                ocr_text = '\n\n'.join(text_results)
                
            except Exception as e:
                print(f"PDF 轉換失敗: {str(e)}")
                return jsonify({
                    "success": False,
                    "error": "PDF 處理失敗，請確認檔案格式正確",
                    "details": str(e)
                }), 500
        
        else:
            # 處理一般圖片
            image = Image.open(file)
            # 指定語言包，包含繁體中文和英文
            ocr_text = pytesseract.image_to_string(
                image, 
                lang='chi_tra+eng',
                config='--psm 1 --oem 1'  # 使用更準確的辨識模式
            )
        
        # 清理辨識出的文字
        ocr_text = ocr_text.strip()
        ocr_text = ocr_text.replace('\n\n', '\n')
        
        if not ocr_text:
            return jsonify({
                "success": True,
                "text": f"由於提供的文字內容無法辨識，無法整理出任何與{CODE_MAPPING.get(code, '未知項目')}相關的條列式重點。",
                "is_raw": True,
                "was_condensed": False
            })

        # 使用 AI 整理文字內容
        summary_prompt = f"""
        請根據以下文字內容，整理成條列式重點，符合以下要求：
        1. 只需列出重要事實與成果，不需要心得反思
        2. 每個重點使用「•」符號開頭
        3. 每個重點不超過30字
        4. 去除與 {CODE_MAPPING.get(code, '未知項目')} 無關的內容
        5. 使用客觀、簡潔的描述方式
        6. 不需要加入個人感想或反思
        7. 不須加入姓名、證號等等資訊

        原文內容：
        {ocr_text}

        請產生條列式重點：
        """

        try:
            gemini_model = genai.GenerativeModel("gemini-2.0-flash")
            gemini_response = gemini_model.generate_content(
                summary_prompt,
                generation_config={"temperature": 0.3}
            )
            summary_text = gemini_response.text.strip()
            
            # 檢查總字數並在需要時進行濃縮
            if len(summary_text) > 300:
                summary_text = condense_text(summary_text, 300)
                return jsonify({
                    "success": True,
                    "text": summary_text,
                    "is_raw": False,
                    "was_condensed": True
                })
            
            return jsonify({
                "success": True,
                "text": summary_text,
                "is_raw": False,
                "was_condensed": False
            })
            
        except Exception as e:
            print(f"AI 摘要處理失敗: {str(e)}")
            # 如果 AI 處理失敗，返回原始辨識文字（但確保不超過300字）
            if len(ocr_text) > 300:
                ocr_text = ocr_text[:300]
            return jsonify({
                "success": True,
                "text": ocr_text,
                "is_raw": True,
                "was_condensed": False
            })
        
    except Exception as e:
        error_msg = f"OCR 處理失敗: {str(e)}"
        print(error_msg)
        return jsonify({
            "success": False,
            "error": error_msg,
            "details": str(e)
        }), 500

# ==== 濃縮文字內容 ====
def condense_text(text, max_length=300):
    if len(text) <= max_length:
        return text
        
    try:
        gemini_model = genai.GenerativeModel("gemini-2.0-flash")
        prompt = f"""
        請將以下文字內容濃縮至{max_length}字以內，保留最重要的資訊：
        
        {text}
        
        要求：
        1. 字數限制：{max_length}字
        2. 保留關鍵事實與重要成果
        3. 去除重複或次要資訊
        4. 保持文意通順
        5. 確保每個檔案都有被提及
        6. 保持條列式格式
        7. 使用「•」符號作為條列點
        """
        
        response = gemini_model.generate_content(
            prompt,
            generation_config={"temperature": 0.3}
        )
        condensed = response.text.strip()
        
        # 如果還是超過字數限制，進行簡單的截斷
        if len(condensed) > max_length:
            # 先嘗試在最後一個完整條列點處截斷
            lines = condensed.split('\n')
            truncated_lines = []
            current_length = 0
            
            for line in lines:
                if current_length + len(line) + 1 <= max_length:  # +1 for newline
                    truncated_lines.append(line)
                    current_length += len(line) + 1
                else:
                    break
            
            if truncated_lines:
                return '\n'.join(truncated_lines)
            else:
                # 如果無法在條列點處截斷，則在句號處截斷
                truncated = condensed[:max_length]
                last_period = truncated.rfind('。')
                if last_period > 0:
                    return truncated[:last_period + 1]
                return truncated
                
        return condensed
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        error_msg = f"濃縮文字時發生錯誤: {str(e)}\n{error_trace}"
        print(error_msg)
        
        # 如果 AI 處理失敗，進行簡單的截斷
        if len(text) > max_length:
            lines = text.split('\n')
            truncated_lines = []
            current_length = 0
            
            for line in lines:
                if current_length + len(line) + 1 <= max_length:
                    truncated_lines.append(line)
                    current_length += len(line) + 1
                else:
                    break
            
            if truncated_lines:
                return '\n'.join(truncated_lines)
            else:
                truncated = text[:max_length]
                last_period = truncated.rfind('。')
                if last_period > 0:
                    return truncated[:last_period + 1]
                return truncated
        return text

# ==== 濃縮文字路由 ====
@app.route("/condense_text", methods=["POST"])
def condense_text_route():
    try:
        data = request.get_json()
        text = data.get('text')
        max_length = data.get('max_length', 300)
        
        if not text:
            return jsonify({
                "success": False, 
                "error": "缺少必要文字內容",
                "details": "請提供需要濃縮的文字"
            }), 400
            
        # 使用 condense_text 函數
        condensed = condense_text(text, max_length)
            
        return jsonify({
            "success": True,
            "text": condensed
        })
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        error_msg = f"處理文字時發生錯誤: {str(e)}"
        print(f"濃縮文字路由錯誤: {error_msg}")
        print(f"錯誤詳情: {error_trace}")
        return jsonify({
            "success": False,
            "error": error_msg,
            "details": error_trace
        }), 500

# ==== 收藏/取消收藏 ====
@app.route("/toggle_favorite", methods=["POST"])
@login_required
def toggle_favorite():
    data = request.get_json()
    history_id = data.get("history_id")
    
    if not history_id:
        return jsonify({"error": "缺少必要參數"}), 400
        
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 確認記錄存在且屬於當前用戶
        cursor.execute("""
            SELECT is_favorite 
            FROM generation_history 
            WHERE id = %s AND user_id = %s
        """, (history_id, current_user.id))
        result = cursor.fetchone()
        
        if not result:
            cursor.close()
            conn.close()
            return jsonify({"error": "找不到記錄或無權限"}), 404
            
        # 切換收藏狀態
        new_status = not result[0]
        cursor.execute("""
            UPDATE generation_history 
            SET is_favorite = %s 
            WHERE id = %s AND user_id = %s
        """, (new_status, history_id, current_user.id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            "success": True,
            "is_favorite": new_status
        })
        
    except Exception as e:
        return jsonify({"error": f"操作失敗: {str(e)}"}), 500

# ==== 從歷史記錄載入內容 ====
@app.route("/load_history", methods=["GET"])
@login_required
def load_history():
    history_id = request.args.get("history_id")
    if not history_id:
        return jsonify({"error": "缺少必要參數"}), 400
        
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 確保只能載入當前用戶的記錄
        cursor.execute("""
            SELECT university, department, section_code, generated_text, style, mindmap_data
            FROM generation_history 
            WHERE id = %s AND user_id = %s
        """, (history_id, current_user.id))
        
        record = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not record:
            return jsonify({"error": "找不到記錄或無權限訪問"}), 404
            
        # 驗證必要欄位
        if not all([record[0], record[1], record[2], record[3]]):  # university, department, section_code, generated_text
            return jsonify({"error": "歷史記錄數據不完整"}), 400
            
        # 安全地解析 mindmap_data
        try:
            mindmap_data = json.loads(record[5]) if record[5] else {}
        except (TypeError, json.JSONDecodeError) as e:
            print(f"解析 mindmap_data 失敗: {str(e)}")
            mindmap_data = {}
            
        response_data = {
            "university": record[0],
            "department": record[1],
            "section_code": record[2],
            "generated_text": record[3],
            "style": record[4],
            "mindmap_data": mindmap_data
        }
            
        return jsonify(response_data)
        
    except Exception as e:
        print(f"載入歷史記錄時發生錯誤: {str(e)}")
        return jsonify({"error": f"載入失敗: {str(e)}"}), 500

# ==== 生成心智圖 ====
@app.route("/generate_mindmap", methods=["POST"])
@login_required
def generate_mindmap():
    data = request.get_json()
    university = data.get("university", "未知大學")
    department = data.get("department", "未知學系")
    section_code = data.get("section_code")
    user_inputs = data.get("user_inputs", {})

    if not section_code or not user_inputs:
        return jsonify({"error": "缺少必要參數"}), 400

    # 獲取學系特色
    department_features = get_department_features(university, department)
    
    # 將使用者輸入轉換為更易讀的格式
    formatted_inputs = "\n".join([
        f"{CODE_MAPPING.get(code, '未知項目')}: {content}" 
        for code, content in user_inputs.items()
    ])

    # 準備心智圖生成提示
    mindmap_prompt = f"""
    請根據以下內容，生成一個心智圖（使用 Mermaid 語法），展示申請者如何透過個人經歷符合 {university}-{department} 的需求：

    【學系特色與需求】
    {department_features}

    【申請項目】
    {section_code} - {CODE_MAPPING.get(section_code, "未知項目")}

    【申請者經歷】
    {formatted_inputs}

    請使用以下規則生成心智圖：
    1. 核心節點應為申請項目名稱
    2. 第一層分支應為2-3個學系所需的關鍵能力或特質
    3. 第二層分支應連結到申請者的具體經歷，說明如何展現或培養這些能力
    4. 所有內容必須基於申請者提供的實際經歷，不得編造或擴充
    5. 用簡潔的文字表達（每個節點建議不超過15字）
    6. 確保每個提到的經歷都能對應到學系要求的能力
    7. 使用中文輸出
    8. 確保每行的縮排使用兩個空格

    請直接輸出 Mermaid mindmap 語法，不要包含任何其他說明文字：

    mindmap
      root((核心主題))
        branch1[學系需求能力1]
          sub1[相關經歷A:心得]
          sub2[相關經歷B:心得]
        branch2[學系需求能力2]
          sub3[相關經歷C:心得]
          sub4[相關經歷D:心得]
    """

    try:
        gemini_model = genai.GenerativeModel("gemini-2.0-flash")
        response = gemini_model.generate_content(
            mindmap_prompt,
            generation_config={
                "temperature": 0.6,
                "max_output_tokens": 1024,
                "top_p": 0.8,
                "top_k": 40
            }
        )
        
        # 提取 Mermaid 語法
        mindmap_text = response.text.strip()
        
        # 如果回應包含 ```mermaid，提取其中的內容
        if "```mermaid" in mindmap_text:
            mindmap_text = mindmap_text.split("```mermaid")[1].split("```")[0].strip()
        
        # 確保文本以 mindmap 開頭
        if not mindmap_text.startswith("mindmap"):
            mindmap_text = "mindmap\n" + mindmap_text
            
        print("生成的心智圖:", mindmap_text)  # 調試用
        
        return jsonify({
            "mindmap_svg": mindmap_text
        })
        
    except Exception as e:
        print(f"生成心智圖時發生錯誤: {str(e)}")  # 調試用
        return jsonify({"error": f"生成心智圖時發生錯誤: {str(e)}"}), 500

def backup_database():
    """備份現有數據庫數據"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 創建備份目錄
        backup_dir = 'database_backups'
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
            
        # 生成備份文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_file = os.path.join(backup_dir, f'backup_{timestamp}.json')
        
        # 備份用戶數據
        cursor.execute("SELECT * FROM users")
        users = cursor.fetchall()
        
        # 備份生成歷史
        cursor.execute("SELECT * FROM generation_history")
        history = cursor.fetchall()
        
        # 保存備份
        backup_data = {
            'users': [dict(zip([column[0] for column in cursor.description], row)) for row in users],
            'generation_history': [dict(zip([column[0] for column in cursor.description], row)) for row in history]
        }
        
        with open(backup_file, 'w', encoding='utf-8') as f:
            json.dump(backup_data, f, ensure_ascii=False, default=str)
            
        cursor.close()
        conn.close()
        return True, backup_file
    except Exception as e:
        print(f"備份失敗: {str(e)}")
        return False, str(e)

def drop_all_tables():
    """刪除所有現有的表"""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        # 先刪除有外鍵依賴的表
        cursor.execute("DROP TABLE IF EXISTS generation_history")
        # 再刪除被依賴的表
        cursor.execute("DROP TABLE IF EXISTS users")
        conn.commit()
    except Exception as e:
        print(f"刪除表失敗: {str(e)}")
        conn.rollback()
    finally:
        cursor.close()
        conn.close()

def init_db():
    """初始化數據庫"""
    print("開始初始化數據庫...")
    
    # 1. 先嘗試備份現有數據
    success, result = backup_database()
    if success:
        print(f"✅ 數據庫備份成功，備份文件: {result}")
    else:
        print(f"⚠️ 數據庫備份失敗: {result}")
        response = input("是否繼續初始化數據庫？(y/n): ")
        if response.lower() != 'y':
            print("❌ 取消初始化操作")
            return
    
    # 2. 刪除現有的表
    print("正在刪除現有的表...")
    drop_all_tables()
    
    # 3. 創建新的表
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 創建用戶表
        cursor.execute("""
            CREATE TABLE users (
                id SERIAL PRIMARY KEY,
                email VARCHAR(255) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,

                -- 登入與驗證
                is_verified BOOLEAN DEFAULT FALSE,
                verification_code VARCHAR(10),
                verification_sent_at TIMESTAMP,

                -- 會員方案
                plan VARCHAR(50) DEFAULT 'Free',
                plan_expiration DATE,

                -- 使用紀錄
                generation_count INTEGER DEFAULT 0,
                generation_quota INTEGER DEFAULT 10,

                -- 時間戳
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        print("✅ 用戶表創建成功")
        
        # 創建生成歷史記錄表
        cursor.execute("""
            CREATE TABLE generation_history (
                id VARCHAR(6) PRIMARY KEY,
                user_id INTEGER REFERENCES users(id),
                university VARCHAR(100) NOT NULL,
                department VARCHAR(100) NOT NULL,
                section_code VARCHAR(10) NOT NULL,
                generated_text TEXT NOT NULL,
                mindmap_data JSONB,
                style VARCHAR(50),
                is_favorite BOOLEAN DEFAULT FALSE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        print("✅ 生成歷史記錄表創建成功")
        
        conn.commit()
        print("✅ 數據庫初始化完成")
        
    except Exception as e:
        print(f"❌ 創建表時發生錯誤: {str(e)}")
        conn.rollback()
    finally:
        cursor.close()
        conn.close()

# 在應用啟動時初始化資料庫
if __name__ == "__main__":
    # 添加命令行參數解析
    import argparse
    parser = argparse.ArgumentParser(description='數據庫管理工具')
    parser.add_argument('--init-db', action='store_true', help='初始化數據庫')
    args = parser.parse_args()
    
    if args.init_db:
        init_db()
    app.run(port=5001, debug=True)

@app.route("/save_generated_content", methods=["POST"])
@login_required
def save_generated_content():
    try:
        data = request.get_json()
        
        # 檢查必要的字段
        required_fields = ['university', 'department', 'section_code', 'generated_text', 'style']
        for field in required_fields:
            if field not in data:
                return jsonify({"success": False, "error": f"Missing required field: {field}"}), 400

        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 從數據庫獲取最新的history_id，確保只查詢當前用戶的記錄
        cursor.execute("""
            SELECT id 
            FROM generation_history 
            WHERE user_id = %s AND university = %s AND department = %s AND section_code = %s
            ORDER BY created_at DESC 
            LIMIT 1
        """, (current_user.id, data['university'], data['department'], data['section_code']))
        
        result = cursor.fetchone()
        
        if result:
            # 更新現有記錄，確保只更新當前用戶的記錄
            history_id = result[0]
            cursor.execute("""
                UPDATE generation_history 
                SET generated_text = %s,
                    mindmap_data = %s,
                    style = %s,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = %s AND user_id = %s
            """, (
                data['generated_text'],
                json.dumps(data.get('mindmap_data', {})),
                data['style'],
                history_id,
                current_user.id
            ))
        else:
            # 生成新的6位亂碼ID
            history_id = generate_random_id()
            
            # 創建新記錄，使用生成的亂碼ID
            cursor.execute("""
                INSERT INTO generation_history (
                    id, user_id, university, department, section_code,
                    generated_text, mindmap_data, style, is_favorite,
                    created_at, updated_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
            """, (
                history_id,
                current_user.id,
                data['university'],
                data['department'],
                data['section_code'],
                data['generated_text'],
                json.dumps(data.get('mindmap_data', {})),
                data['style'],
                False
            ))

        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({"success": True, "history_id": history_id})

    except Exception as e:
        print("Error saving content:", str(e))
        return jsonify({"success": False, "error": str(e)}), 500

def generate_random_id():
    """生成6位亂碼ID，包含數字和大寫字母"""
    characters = string.ascii_uppercase + string.digits  # A-Z and 0-9
    while True:
        # 生成6位亂碼
        random_id = ''.join(random.choices(characters, k=6))
        
        # 檢查ID是否已存在
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM generation_history WHERE id = %s", (random_id,))
        exists = cursor.fetchone()
        cursor.close()
        conn.close()
        
        # 如果ID不存在，則返回這個ID
        if not exists:
            return random_id
